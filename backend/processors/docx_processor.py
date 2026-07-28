"""
docx_processor.py — DOCX Text Extraction Engine
================================================
Handles Microsoft Word .docx file text extraction using python-docx.

Key difference from PDF processing:
  - DOCX has semantic structure: paragraphs, headings, tables
  - We extract structure-aware text, not just raw characters
  - Tables are converted to readable text (not silently dropped)
  - Heading levels are preserved as markers for better chunking

Design: Same "never raise" contract as pdf_processor.py.
        Always returns a ProcessedDocument, errors included in it.
"""

import time
from pathlib import Path
from typing import Optional

# python-docx is the main library for DOCX parsing
# It understands the underlying XML structure of .docx files
import docx
from docx import Document
from docx.oxml.ns import qn          # XML namespace helper
from docx.table import Table          # Table type for isinstance checks
from docx.text.paragraph import Paragraph  # Paragraph type

from backend.models.document_models import (
    DocumentMetadata,
    DocumentType,
    PageContent,
    ProcessedDocument,
    ProcessingStatus,
)
from backend.utils.logger import logger


# ── Constants ──────────────────────────────────────────────────────
MAX_FILE_SIZE_BYTES = 100 * 1024 * 1024   # 100 MB limit

# DOCX heading styles — we use these to add structure markers
# python-docx uses "Heading 1", "Heading 2", etc. as style names
HEADING_STYLES = {
    "Heading 1": "##",     # Markdown-style markers for importance
    "Heading 2": "###",
    "Heading 3": "####",
    "Title": "#",
    "Subtitle": "##",
}


class DOCXProcessor:
    """
    Extracts text from Microsoft Word .docx files.

    Handles:
      - Regular paragraphs and headings
      - Tables (converts rows to tab-separated text)
      - Documents with only images (returns EMPTY status)
      - Corrupted or invalid .docx files (returns FAILED status)

    Usage:
        processor = DOCXProcessor()
        result = processor.process("path/to/file.docx")
        print(result.full_text)
    """

    def __init__(self):
        logger.debug("DOCXProcessor initialized")

    def process(self, file_path: str | Path) -> ProcessedDocument:
        """
        Main entry point. Process a DOCX file end-to-end.

        Args:
            file_path: Path to the .docx file

        Returns:
            ProcessedDocument — never raises, errors in result object.
        """
        start_time = time.time()
        file_path = Path(file_path)

        logger.info(f"Starting DOCX processing: {file_path.name}")

        # ── Step 1: Validate ───────────────────────────────────────
        validation_error = self._validate_file(file_path)
        if validation_error:
            validation_error.processing_time = time.time() - start_time
            return validation_error

        # ── Step 2: Build initial metadata ────────────────────────
        file_size = file_path.stat().st_size
        metadata = DocumentMetadata(
            filename=file_path.name,
            file_size_bytes=file_size,
            document_type=DocumentType.DOCX,
            page_count=0,   # DOCX doesn't have true pages — will estimate
        )

        # ── Step 3: Open the document ──────────────────────────────
        try:
            doc = Document(str(file_path))
        except Exception as e:
            logger.error(f"Failed to open DOCX {file_path.name}: {e}")
            return ProcessedDocument(
                metadata=metadata,
                pages=[],
                full_text="",
                status=ProcessingStatus.FAILED,
                error_message=(
                    f"Could not open DOCX file. "
                    f"It may be corrupted or not a valid Word document: {e}"
                ),
                processing_time=time.time() - start_time,
            )

        # ── Step 4: Extract document-level metadata ────────────────
        self._extract_docx_metadata(doc, metadata)

        # ── Step 5: Extract all content (paragraphs + tables) ─────
        # DOCX body is a sequence of Block elements.
        # Blocks are either Paragraph or Table objects — in document order.
        content_blocks = self._extract_content_blocks(doc, file_path.name)

        # ── Step 6: Build the full text string ────────────────────
        full_text = self._assemble_full_text(content_blocks)

        # ── Step 7: Estimate "pages" from word count ──────────────
        # DOCX has no true page concept — we estimate.
        # Average page = ~250 words. This is used for metadata only.
        word_count = len(full_text.split())
        estimated_pages = max(1, round(word_count / 250))
        metadata.page_count = estimated_pages
        metadata.total_chars = len(full_text)
        metadata.total_words = word_count

        # ── Step 8: Create synthetic "pages" for the output model ──
        # Our ProcessedDocument expects a list of PageContent.
        # For DOCX, we create section-based "pages" rather than true pages.
        pages = self._create_section_pages(content_blocks, estimated_pages)

        # ── Step 9: Determine status ───────────────────────────────
        status, error_message = self._determine_status(full_text, file_path.name)

        processing_time = time.time() - start_time

        logger.info(
            f"DOCX processed: {file_path.name} | "
            f"Status: {status} | "
            f"~{estimated_pages} pages | "
            f"Chars: {metadata.total_chars} | "
            f"Words: {metadata.total_words} | "
            f"Time: {processing_time:.2f}s"
        )

        return ProcessedDocument(
            metadata=metadata,
            pages=pages,
            full_text=full_text,
            status=status,
            error_message=error_message,
            processing_time=processing_time,
        )

    # ── Private Methods ────────────────────────────────────────────

    def _validate_file(self, file_path: Path) -> Optional[ProcessedDocument]:
        """
        Validate the DOCX file before processing.
        Same pattern as PDFProcessor._validate_file().
        """
        if not file_path.exists():
            return self._make_error_doc(
                filename=file_path.name,
                status=ProcessingStatus.FAILED,
                error=f"File not found: {file_path}",
            )

        if not file_path.is_file():
            return self._make_error_doc(
                filename=file_path.name,
                status=ProcessingStatus.FAILED,
                error=f"Path is not a file: {file_path}",
            )

        file_size = file_path.stat().st_size

        if file_size == 0:
            return self._make_error_doc(
                filename=file_path.name,
                status=ProcessingStatus.FAILED,
                error="File is empty (0 bytes)",
            )

        if file_size > MAX_FILE_SIZE_BYTES:
            size_mb = file_size / (1024 * 1024)
            return self._make_error_doc(
                filename=file_path.name,
                status=ProcessingStatus.FAILED,
                error=f"File too large: {size_mb:.1f}MB. Maximum: 100MB",
            )

        # .docx check — also accept .doc? No: .doc is old binary format,
        # not supported by python-docx. We reject it clearly.
        suffix = file_path.suffix.lower()
        if suffix == ".doc":
            return self._make_error_doc(
                filename=file_path.name,
                status=ProcessingStatus.FAILED,
                error=(
                    "Old .doc format is not supported. "
                    "Please save as .docx in Microsoft Word and re-upload."
                ),
            )

        if suffix != ".docx":
            return self._make_error_doc(
                filename=file_path.name,
                status=ProcessingStatus.FAILED,
                error=f"Not a DOCX file. Extension found: {suffix}",
            )

        return None  # All checks passed

    def _extract_docx_metadata(
        self, doc: Document, metadata: DocumentMetadata
    ) -> None:
        """
        Extract document properties (title, author, etc.) from DOCX.

        DOCX stores metadata in doc.core_properties — a standard
        Office Open XML (OOXML) properties block.
        """
        try:
            props = doc.core_properties

            # Each property can raise AttributeError if not set
            metadata.title = getattr(props, "title", None) or None
            metadata.author = getattr(props, "author", None) or None
            metadata.subject = getattr(props, "subject", None) or None

            # Creator is 'last_modified_by' in DOCX terms
            metadata.creator = getattr(props, "last_modified_by", None) or None

        except Exception as e:
            logger.warning(f"Could not extract DOCX metadata: {e}")

    def _extract_content_blocks(
        self, doc: Document, filename: str
    ) -> list[dict]:
        """
        Extract all content from the DOCX in document order.

        DOCX body contains a mix of Paragraph and Table elements.
        We iterate them in order to preserve document structure.

        Returns a list of dicts:
            {
                "type": "paragraph" | "table" | "heading",
                "text": str,
                "level": int (for headings, 1-6; 0 for paragraphs),
            }
        """
        blocks = []

        # doc.element.body is the raw XML body element.
        # Iterating it gives us child elements in document order.
        # Each child is either a paragraph (w:p) or table (w:tbl).
        body = doc.element.body

        # We iterate over body children and wrap them in python-docx objects
        for child in body.iterchildren():
            # ── Paragraph (includes headings) ──────────────────────
            # XML tag for paragraph is {http://...}p
            # qn() converts "w:p" → the full qualified XML name
            if child.tag == qn("w:p"):
                paragraph = Paragraph(child, doc)
                block = self._process_paragraph(paragraph)
                if block:  # Skip empty paragraphs
                    blocks.append(block)

            # ── Table ──────────────────────────────────────────────
            elif child.tag == qn("w:tbl"):
                table = Table(child, doc)
                block = self._process_table(table, filename)
                if block:
                    blocks.append(block)

            # Other element types (images, shapes, etc.) are skipped
            # They produce no text content

        logger.debug(
            f"{filename}: extracted {len(blocks)} content blocks"
        )
        return blocks

    def _process_paragraph(self, paragraph: Paragraph) -> Optional[dict]:
        """
        Convert a single DOCX paragraph to a content block.

        Determines if it's a heading or regular paragraph,
        and prefixes headings with markers.

        Args:
            paragraph: python-docx Paragraph object

        Returns:
            Dict with type, text, level — or None if paragraph is empty
        """
        # paragraph.text joins all runs (text segments) in the paragraph
        text = paragraph.text.strip()

        if not text:
            return None  # Skip blank paragraphs

        # Check paragraph style to identify headings
        style_name = paragraph.style.name if paragraph.style else ""

        if style_name in HEADING_STYLES:
            # It's a heading — add marker for semantic importance
            marker = HEADING_STYLES[style_name]
            # Extract heading level number (e.g., "Heading 2" → 2)
            level = self._get_heading_level(style_name)
            return {
                "type": "heading",
                "text": f"{marker} {text}",
                "level": level,
            }
        else:
            return {
                "type": "paragraph",
                "text": text,
                "level": 0,
            }

    def _get_heading_level(self, style_name: str) -> int:
        """
        Extract numeric level from heading style name.

        "Heading 1" → 1
        "Heading 2" → 2
        "Title"     → 0
        """
        if "Heading" in style_name:
            try:
                return int(style_name.split()[-1])
            except (ValueError, IndexError):
                return 1
        return 0

    def _process_table(
        self, table: Table, filename: str
    ) -> Optional[dict]:
        """
        Convert a DOCX table to readable text.

        Tables are structured data. We convert each row to a
        pipe-separated string — similar to Markdown table format.
        This preserves the data for the LLM to understand.

        Example:
            | Name | Age | City      |
            | John | 30  | Jaipur    |
            | Jane | 25  | Delhi     |

        Args:
            table:    python-docx Table object
            filename: For logging only

        Returns:
            Dict with type="table" and text containing all rows
        """
        try:
            rows_text = []

            for row_index, row in enumerate(table.rows):
                # Extract text from each cell, stripping whitespace
                cells = [cell.text.strip() for cell in row.cells]

                # Skip completely empty rows
                if not any(cells):
                    continue

                # Format as pipe-separated (readable and LLM-friendly)
                row_text = " | ".join(cells)
                rows_text.append(f"| {row_text} |")

                # Add separator after header row (row 0)
                if row_index == 0:
                    separator = "|" + "|".join(
                        [" --- " for _ in cells]
                    ) + "|"
                    rows_text.append(separator)

            if not rows_text:
                return None

            table_text = "\n".join(rows_text)
            return {
                "type": "table",
                "text": f"[Table]\n{table_text}",
                "level": 0,
            }

        except Exception as e:
            logger.warning(f"{filename}: table extraction failed: {e}")
            return None

    def _assemble_full_text(self, content_blocks: list[dict]) -> str:
        """
        Join all content blocks into a single readable string.

        Uses double newline between blocks for clear separation.
        Headings get extra spacing above them.
        """
        parts = []

        for block in content_blocks:
            text = block["text"]

            # Add extra blank line before headings for visual separation
            if block["type"] == "heading":
                parts.append(f"\n{text}")
            else:
                parts.append(text)

        return "\n\n".join(parts).strip()

    def _create_section_pages(
        self, content_blocks: list[dict], estimated_pages: int
    ) -> list[PageContent]:
        """
        Create synthetic PageContent objects for DOCX.

        Since DOCX has no real pages, we split the content into
        roughly equal "sections" based on heading boundaries.
        This gives the rest of the pipeline consistent PageContent objects.

        Strategy:
          - Try to split at heading boundaries
          - If no headings, split evenly by word count
        """
        if not content_blocks:
            return []

        # Find heading positions (natural section boundaries)
        heading_positions = [
            i for i, b in enumerate(content_blocks)
            if b["type"] == "heading" and b["level"] <= 2
        ]

        if heading_positions:
            # Split at heading boundaries
            sections = self._split_at_headings(
                content_blocks, heading_positions
            )
        else:
            # No headings — create one section per estimated page
            sections = self._split_evenly(content_blocks, estimated_pages)

        # Convert sections to PageContent objects
        pages = []
        for idx, section_blocks in enumerate(sections, start=1):
            section_text = "\n\n".join(b["text"] for b in section_blocks)
            pages.append(PageContent(
                page_number=idx,
                text=section_text,
            ))

        return pages

    def _split_at_headings(
        self,
        blocks: list[dict],
        heading_positions: list[int],
    ) -> list[list[dict]]:
        """Split content blocks into sections at each heading."""
        sections = []
        positions = [0] + heading_positions + [len(blocks)]

        for i in range(len(positions) - 1):
            start = positions[i]
            end = positions[i + 1]
            section = blocks[start:end]
            if section:
                sections.append(section)

        return sections

    def _split_evenly(
        self, blocks: list[dict], num_sections: int
    ) -> list[list[dict]]:
        """Split blocks into roughly equal sections."""
        if num_sections <= 1 or len(blocks) <= num_sections:
            return [blocks]

        chunk_size = max(1, len(blocks) // num_sections)
        return [
            blocks[i:i + chunk_size]
            for i in range(0, len(blocks), chunk_size)
        ]

    def _determine_status(
        self, full_text: str, filename: str
    ) -> tuple[ProcessingStatus, Optional[str]]:
        """Determine processing status from the assembled text."""

        if not full_text or len(full_text.strip()) == 0:
            return (
                ProcessingStatus.EMPTY,
                (
                    f"No text could be extracted from {filename}. "
                    "The document may contain only images or shapes."
                ),
            )

        if len(full_text.strip()) < 50:
            return (
                ProcessingStatus.PARTIAL,
                f"Very little text extracted ({len(full_text)} chars). "
                "Document may be mostly images.",
            )

        return ProcessingStatus.SUCCESS, None

    def _make_error_doc(
        self,
        filename: str,
        status: ProcessingStatus,
        error: str,
    ) -> ProcessedDocument:
        """Helper to create error ProcessedDocument. Mirrors PDFProcessor."""
        logger.error(f"DOCX validation failed for {filename}: {error}")
        return ProcessedDocument(
            metadata=DocumentMetadata(
                filename=filename,
                file_size_bytes=0,
                document_type=DocumentType.DOCX,
                page_count=0,
            ),
            pages=[],
            full_text="",
            status=status,
            error_message=error,
        )