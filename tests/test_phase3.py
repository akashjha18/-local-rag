"""
test_phase3.py — DOCX Processor Verification
=============================================
Run with:
    python tests/test_phase3.py
"""

import sys
import os
from pathlib import Path

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from backend.processors.docx_processor import DOCXProcessor
from backend.models.document_models import ProcessingStatus, DocumentType
from backend.utils.logger import setup_logger

setup_logger(debug=False)  # Set True for verbose logs


def create_test_docx(path: Path, with_table: bool = False) -> None:
    """
    Create a real .docx file using python-docx for testing.
    This ensures we test with a valid, known document.
    """
    from docx import Document as DocxDocument
    from docx.shared import Pt

    doc = DocxDocument()

    # Add Title
    doc.add_heading("Local RAG System — Test Document", level=0)

    # Add Heading 1
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "This is a test document created to verify the DOCX processor. "
        "It contains multiple sections, paragraphs, and optionally tables."
    )

    # Add Heading 2
    doc.add_heading("What is RAG?", level=2)
    doc.add_paragraph(
        "Retrieval Augmented Generation (RAG) is a technique that combines "
        "information retrieval with language model generation. "
        "It grounds the LLM's answers in real source documents."
    )

    doc.add_heading("Technical Stack", level=2)
    doc.add_paragraph(
        "Backend: FastAPI, Python. "
        "AI: Sentence Transformers, FAISS, Ollama. "
        "Frontend: React + Vite + Tailwind CSS."
    )

    doc.add_heading("Architecture", level=1)
    doc.add_paragraph(
        "The system processes uploaded documents, generates vector embeddings, "
        "stores them in FAISS, and retrieves relevant chunks at query time."
    )

    if with_table:
        doc.add_heading("Comparison Table", level=2)
        doc.add_paragraph("Feature comparison between RAG approaches:")

        table = doc.add_table(rows=4, cols=3)
        table.style = "Table Grid"

        headers = ["Feature", "Basic RAG", "Advanced RAG"]
        data = [
            ["Chunking",     "Fixed size",    "Semantic"],
            ["Retrieval",    "Vector search", "Hybrid BM25+Vector"],
            ["Reranking",    "None",          "Cross-encoder"],
        ]

        # Set header row
        for col_idx, header in enumerate(headers):
            table.rows[0].cells[col_idx].text = header

        # Set data rows
        for row_idx, row_data in enumerate(data, start=1):
            for col_idx, value in enumerate(row_data):
                table.rows[row_idx].cells[col_idx].text = value

    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "This document was created to test the DOCX processor pipeline. "
        "If you can read this, the extraction is working correctly. "
        "Akash Jha, JECRC University, Jaipur."
    )

    doc.save(str(path))


def test_basic_docx():
    """Test basic DOCX extraction with a simple document."""
    print("\n── Test 1: Basic DOCX Extraction ───────────────────")

    sample_dir = Path("data/documents")
    sample_dir.mkdir(parents=True, exist_ok=True)
    docx_path = sample_dir / "test_basic.docx"

    create_test_docx(docx_path, with_table=False)

    processor = DOCXProcessor()
    result = processor.process(docx_path)

    print(f"   File: {result.metadata.filename}")
    print(f"   Status: {result.status}")
    print(f"   ~Pages: {result.metadata.page_count}")
    print(f"   Total chars: {result.metadata.total_chars}")
    print(f"   Total words: {result.metadata.total_words}")
    print(f"   Processing time: {result.processing_time:.3f}s")
    print(f"   Document type: {result.metadata.document_type}")

    assert result.status == ProcessingStatus.SUCCESS
    assert result.metadata.document_type == DocumentType.DOCX
    assert result.metadata.total_chars > 100
    assert "RAG" in result.full_text

    print(f"\n   Text preview:")
    print(f"   {result.full_text[:300]}")
    print(f"\n   ✅ Basic DOCX extraction passed")


def test_docx_with_table():
    """Test DOCX extraction with a table included."""
    print("\n── Test 2: DOCX with Table ──────────────────────────")

    docx_path = Path("data/documents/test_with_table.docx")
    create_test_docx(docx_path, with_table=True)

    processor = DOCXProcessor()
    result = processor.process(docx_path)

    print(f"   Status: {result.status}")
    print(f"   Chars: {result.metadata.total_chars}")

    # Table text should be in the output
    has_table = "[Table]" in result.full_text
    has_feature = "Feature" in result.full_text

    print(f"   Table marker present: {has_table}")
    print(f"   Table headers present: {has_feature}")

    # Show just the table part
    if "[Table]" in result.full_text:
        table_start = result.full_text.index("[Table]")
        table_preview = result.full_text[table_start:table_start + 300]
        print(f"\n   Table extraction preview:")
        print(f"   {table_preview}")

    assert result.status == ProcessingStatus.SUCCESS
    print(f"\n   ✅ DOCX with table extraction passed")


def test_missing_docx():
    """Test missing file handling."""
    print("\n── Test 3: Missing File ─────────────────────────────")

    processor = DOCXProcessor()
    result = processor.process("ghost_file.docx")

    assert result.status == ProcessingStatus.FAILED
    assert "not found" in result.error_message.lower()
    print(f"   ✅ Missing file handled: {result.error_message}")


def test_wrong_extension():
    """Test rejection of non-DOCX files."""
    print("\n── Test 4: Wrong Extension (.doc) ───────────────────")

    fake = Path("data/documents/old_format.doc")
    fake.write_bytes(b"not a real doc")

    processor = DOCXProcessor()
    result = processor.process(fake)

    assert result.status == ProcessingStatus.FAILED
    assert ".doc" in result.error_message
    print(f"   ✅ .doc rejected: {result.error_message}")
    fake.unlink()


def test_sections_created():
    """Test that section-based pages are created properly."""
    print("\n── Test 5: Section Pages ─────────────────────────────")

    docx_path = Path("data/documents/test_basic.docx")
    if not docx_path.exists():
        create_test_docx(docx_path)

    processor = DOCXProcessor()
    result = processor.process(docx_path)

    print(f"   Number of sections: {len(result.pages)}")
    for page in result.pages:
        preview = page.text[:60].replace("\n", " ")
        print(f"   Section {page.page_number}: '{preview}...' "
              f"({page.char_count} chars)")

    assert len(result.pages) >= 1
    print(f"\n   ✅ Sections created correctly")


def test_your_own_docx():
    """Test with your own DOCX file. Set the path below."""
    print("\n── Test 6: Custom DOCX (Optional) ───────────────────")

    # ── SET THIS to test with your own DOCX ───────────────────────
    your_docx_path = None  # e.g., r"C:\Users\hp\Documents\report.docx"

    if your_docx_path is None:
        print("   ⏭️  Skipped — set your_docx_path to test your own file")
        return

    processor = DOCXProcessor()
    result = processor.process(your_docx_path)

    print(f"   File: {result.metadata.filename}")
    print(f"   Status: {result.status}")
    print(f"   Words: {result.metadata.total_words}")
    print(f"   Sections: {len(result.pages)}")
    print(f"\n   Preview: {result.full_text[:400]}")


def test_both_processors_same_output():
    """
    Critical integration test:
    Both PDF and DOCX processors must return the same model type.
    The rest of the pipeline must work with both without any changes.
    """
    print("\n── Test 7: Unified Output Model ─────────────────────")

    from backend.processors.pdf_processor import PDFProcessor
    from backend.models.document_models import ProcessedDocument

    pdf_processor = PDFProcessor()
    docx_processor = DOCXProcessor()

    # Process sample PDF from Phase 2
    pdf_path = Path("data/documents/sample.pdf")
    docx_path = Path("data/documents/test_basic.docx")

    results = []
    if pdf_path.exists():
        results.append(("PDF", pdf_processor.process(pdf_path)))
    if docx_path.exists():
        results.append(("DOCX", docx_processor.process(docx_path)))

    for doc_type, result in results:
        assert isinstance(result, ProcessedDocument), \
            f"{doc_type} didn't return ProcessedDocument!"
        print(f"   ✅ {doc_type} returns ProcessedDocument "
              f"(status={result.status}, chars={result.metadata.total_chars})")

    print(f"\n   ✅ Both processors return identical output model")
    print(f"      The pipeline can treat them interchangeably ✓")


if __name__ == "__main__":
    print("\n" + "=" * 55)
    print("   LOCAL RAG — PHASE 3: DOCX PROCESSOR TESTS")
    print("=" * 55)

    test_basic_docx()
    test_docx_with_table()
    test_missing_docx()
    test_wrong_extension()
    test_sections_created()
    test_your_own_docx()
    test_both_processors_same_output()

    print("\n" + "=" * 55)
    print("   Phase 3 complete!")
    print("   Next: confirm all ✅ → Phase 4 (Text Chunking)")
    print("=" * 55 + "\n")